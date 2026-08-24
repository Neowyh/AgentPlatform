"""Tests for ideer.utils.docx_rich (rich Word -> Markdown parsing).

Builds real .docx fixtures programmatically with python-docx and verifies the
end-to-end conversion: headings, inline formatting, lists, tables, image
artifact relocation, and legacy .doc guard rails.
"""

from __future__ import annotations

import pytest

from ideer.utils.docx_rich import convert_docx, is_available, libreoffice_available

pytestmark = pytest.mark.skipif(not is_available(), reason="python-docx is not installed")


def _build_sample_docx(path):
    """Create a small but representative .docx fixture."""
    from docx import Document

    doc = Document()

    doc.add_paragraph("项目计划", style="Heading 1")
    doc.add_paragraph("概述", style="Heading 2")

    para = doc.add_paragraph()
    run = para.add_run("重要")
    run.bold = True
    para.add_run("：本项目需要按时交付。")

    doc.add_paragraph("第一步", style="List Bullet")
    doc.add_paragraph("第二步", style="List Bullet")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "名称"
    table.rows[0].cells[1].text = "数值"
    table.rows[1].cells[0].text = "吞吐"
    table.rows[1].cells[1].text = "100"

    doc.save(str(path))


class TestConvertDocx:
    def test_headings_and_formatting(self, tmp_path):
        src = tmp_path / "plan.docx"
        _build_sample_docx(src)

        markdown, extra_info = convert_docx(src, tmp_path)

        assert "# 项目计划" in markdown
        assert "## 概述" in markdown
        # Bold formatting preserved
        assert "**重要**" in markdown
        # Bullet lists preserved
        assert "- 第一步" in markdown
        assert "- 第二步" in markdown
        assert "output_subdir" in extra_info
        assert "image_dir" in extra_info

    def test_tables_rendered_as_html(self, tmp_path):
        src = tmp_path / "table.docx"
        _build_sample_docx(src)

        markdown, _ = convert_docx(src, tmp_path)

        assert "<table>" in markdown
        assert "<th>名称</th>" in markdown
        assert "<td>100</td>" in markdown

    def test_artifact_layout(self, tmp_path):
        """Parsed artifacts live under <stem>_<ext>/ (relocation to
        <stem>_files/ happens later in file_conversion._convert_word_with_rich)."""
        src = tmp_path / "layout.docx"
        _build_sample_docx(src)

        convert_docx(src, tmp_path)

        assert (tmp_path / "layout_docx").exists()
        assert (tmp_path / "layout_docx" / "images").exists()

    def test_markdown_written_to_output_dir(self, tmp_path):
        """The parser writes an intermediate <stem>.md inside its subdirectory."""
        src = tmp_path / "saved.docx"
        _build_sample_docx(src)

        convert_docx(src, tmp_path)

        inner_md = tmp_path / "saved_docx" / "saved.md"
        assert inner_md.exists()


class TestLegacyDocGuard:
    def test_doc_without_libreoffice_raises(self, tmp_path):
        """Legacy .doc requires soffice; a clear error is raised without it."""
        if libreoffice_available():
            pytest.skip("LibreOffice is installed on this host")

        doc = tmp_path / "legacy.doc"
        doc.write_bytes(b"\xd0\xcf\x11\xe0 fake OLE")  # fake .doc bytes

        with pytest.raises(RuntimeError, match="LibreOffice"):
            convert_docx(doc, tmp_path)

    def test_unsupported_extension_raises(self, tmp_path):
        txt = tmp_path / "notes.txt"
        txt.write_text("hello", encoding="utf-8")

        with pytest.raises(ValueError, match="Unsupported file format"):
            convert_docx(txt, tmp_path)


class TestParseDocTempCleanup:
    def test_converted_docx_temp_dir_removed(self, tmp_path, monkeypatch):
        """.doc 转换用的临时目录在解析完成后必须被清理。"""
        from ideer.utils.docx_rich.parser import DocxDocumentParser

        src = tmp_path / "legacy.doc"
        src.write_bytes(b"\xd0\xcf\x11\xe0 fake OLE")

        # 预先准备一个真实 docx，伪装成 LibreOffice 的转换产物
        convert_dir = tmp_path / "_converted_docx"
        convert_dir.mkdir()
        converted = convert_dir / "legacy.docx"
        _build_sample_docx(converted)

        class FakeLibre:
            def convert_to_modern_format(self, file_path, output_dir, timeout=60):
                return str(converted)

        monkeypatch.setattr(
            "ideer.utils.docx_rich.libreoffice.get_libreoffice_manager",
            lambda: FakeLibre(),
        )

        content, _, _extra = DocxDocumentParser().parse_doc(str(src), str(tmp_path))

        assert "# 项目计划" in content
        assert not convert_dir.exists()
