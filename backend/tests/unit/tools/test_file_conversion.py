"""Tests for file_conversion utilities (PR1: pymupdf4llm + asyncio.to_thread; PR2: extract_outline)."""

from __future__ import annotations

import asyncio
import sys
from types import ModuleType
from unittest.mock import MagicMock, patch

import pytest

from ideer.utils.file_conversion import (
    _ASYNC_THRESHOLD_BYTES,
    _MIN_CHARS_PER_PAGE,
    MAX_OUTLINE_ENTRIES,
    _do_convert,
    _get_docx_converter,
    _get_pdf_converter,
    _pymupdf_output_too_sparse,
    convert_file_to_markdown,
    extract_outline,
)


def _make_pymupdf_mock(page_count: int) -> ModuleType:
    """Return a fake *pymupdf* module whose ``open()`` reports *page_count* pages."""
    mock_doc = MagicMock()
    mock_doc.__len__ = MagicMock(return_value=page_count)
    fake_pymupdf = ModuleType("pymupdf")
    fake_pymupdf.open = MagicMock(return_value=mock_doc)  # type: ignore[attr-defined]
    return fake_pymupdf


def _run(coro):
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


# ---------------------------------------------------------------------------
# _pymupdf_output_too_sparse
# ---------------------------------------------------------------------------


class TestPymupdfOutputTooSparse:
    """Check the chars-per-page sparsity heuristic."""

    def test_dense_text_pdf_not_sparse(self, tmp_path):
        """Normal text PDF: many chars per page → not sparse."""
        pdf = tmp_path / "dense.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        # 10 pages × 10 000 chars → 1000/page ≫ threshold
        with patch.dict(sys.modules, {"pymupdf": _make_pymupdf_mock(page_count=10)}):
            result = _pymupdf_output_too_sparse("x" * 10_000, pdf)
        assert result is False

    def test_image_based_pdf_is_sparse(self, tmp_path):
        """Image-based PDF: near-zero chars per page → sparse."""
        pdf = tmp_path / "image.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        # 612 chars / 31 pages ≈ 19.7/page < _MIN_CHARS_PER_PAGE (50)
        with patch.dict(sys.modules, {"pymupdf": _make_pymupdf_mock(page_count=31)}):
            result = _pymupdf_output_too_sparse("x" * 612, pdf)
        assert result is True

    def test_fallback_when_pymupdf_unavailable(self, tmp_path):
        """When pymupdf is not installed, fall back to absolute 200-char threshold."""
        pdf = tmp_path / "broken.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        # Remove pymupdf from sys.modules so the `import pymupdf` inside the
        # function raises ImportError, triggering the absolute-threshold fallback.
        with patch.dict(sys.modules, {"pymupdf": None}):
            sparse = _pymupdf_output_too_sparse("x" * 100, pdf)
            not_sparse = _pymupdf_output_too_sparse("x" * 300, pdf)

        assert sparse is True
        assert not_sparse is False

    def test_exactly_at_threshold_is_not_sparse(self, tmp_path):
        """Chars-per-page == threshold is treated as NOT sparse (boundary inclusive)."""
        pdf = tmp_path / "boundary.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        # 2 pages × _MIN_CHARS_PER_PAGE chars = exactly at threshold
        with patch.dict(sys.modules, {"pymupdf": _make_pymupdf_mock(page_count=2)}):
            result = _pymupdf_output_too_sparse("x" * (_MIN_CHARS_PER_PAGE * 2), pdf)
        assert result is False


# ---------------------------------------------------------------------------
# _do_convert — routing logic
# ---------------------------------------------------------------------------


class TestDoConvert:
    """Verify that _do_convert routes to the right sub-converter."""

    def test_xlsx_pptx_always_use_markitdown(self, tmp_path):
        """XLSX / PPTX always go through MarkItDown regardless of setting."""
        xlsx = tmp_path / "book.xlsx"
        xlsx.write_bytes(b"PK fake xlsx")

        with (
            patch("ideer.utils.file_conversion._convert_word_with_rich") as mock_rich,
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="# Markdown from MarkItDown",
            ) as mock_md,
        ):
            result = _do_convert(xlsx, "auto", "auto")

        mock_rich.assert_not_called()
        mock_md.assert_called_once_with(xlsx)
        assert result == "# Markdown from MarkItDown"

    def test_docx_markitdown_mode_skips_rich(self, tmp_path):
        """docx_converter='markitdown': legacy behaviour, rich parser never invoked."""
        docx = tmp_path / "report.docx"
        docx.write_bytes(b"PK fake docx")

        with (
            patch("ideer.utils.file_conversion._convert_word_with_rich") as mock_rich,
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="# Markdown from MarkItDown",
            ) as mock_md,
        ):
            result = _do_convert(docx, "auto", "markitdown")

        mock_rich.assert_not_called()
        mock_md.assert_called_once_with(docx)
        assert result == "# Markdown from MarkItDown"

    def test_docx_auto_prefers_rich(self, tmp_path):
        """auto mode: rich Word parser wins when it succeeds."""
        docx = tmp_path / "report.docx"
        docx.write_bytes(b"PK fake docx")

        with (
            patch(
                "ideer.utils.file_conversion._convert_word_with_rich",
                return_value="# Rich Heading\n",
            ) as mock_rich,
            patch("ideer.utils.file_conversion._convert_with_markitdown") as mock_md,
        ):
            result = _do_convert(docx, "auto", "auto")

        mock_rich.assert_called_once_with(docx)
        mock_md.assert_not_called()
        assert result == "# Rich Heading\n"

    def test_docx_auto_falls_back_when_rich_unavailable(self, tmp_path):
        """auto mode: fall back to MarkItDown when rich parsing is unavailable."""
        docx = tmp_path / "report.docx"
        docx.write_bytes(b"PK fake docx")

        with (
            patch(
                "ideer.utils.file_conversion._convert_word_with_rich",
                return_value=None,
            ),
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="MarkItDown fallback",
            ) as mock_md,
        ):
            result = _do_convert(docx, "auto", "auto")

        mock_md.assert_called_once_with(docx)
        assert result == "MarkItDown fallback"

    def test_docx_rich_mode_raises_when_unavailable(self, tmp_path):
        """'rich' mode: failure surfaces instead of silently degrading."""
        docx = tmp_path / "report.docx"
        docx.write_bytes(b"PK fake docx")

        with (
            patch(
                "ideer.utils.file_conversion._convert_word_with_rich",
                return_value=None,
            ),
            patch("ideer.utils.file_conversion._convert_with_markitdown") as mock_md,
        ):
            try:
                _do_convert(docx, "auto", "rich")
                raised = False
            except RuntimeError:
                raised = True

        assert raised
        mock_md.assert_not_called()

    def test_pdf_auto_uses_pymupdf4llm_when_dense(self, tmp_path):
        """auto mode: use pymupdf4llm output when it's dense enough."""
        pdf = tmp_path / "report.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        dense_text = "# Heading\n" + "word " * 2000  # clearly dense

        with (
            patch(
                "ideer.utils.file_conversion._convert_pdf_with_pymupdf4llm",
                return_value=dense_text,
            ),
            patch(
                "ideer.utils.file_conversion._pymupdf_output_too_sparse",
                return_value=False,
            ),
            patch("ideer.utils.file_conversion._convert_with_markitdown") as mock_md,
        ):
            result = _do_convert(pdf, "auto")

        mock_md.assert_not_called()
        assert result == dense_text

    def test_pdf_auto_falls_back_when_sparse(self, tmp_path):
        """auto mode: fall back to MarkItDown when pymupdf4llm output is sparse."""
        pdf = tmp_path / "scanned.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        with (
            patch(
                "ideer.utils.file_conversion._convert_pdf_with_pymupdf4llm",
                return_value="x" * 612,  # 19.7 chars/page for 31-page doc
            ),
            patch(
                "ideer.utils.file_conversion._pymupdf_output_too_sparse",
                return_value=True,
            ),
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="OCR result via MarkItDown",
            ) as mock_md,
        ):
            result = _do_convert(pdf, "auto")

        mock_md.assert_called_once_with(pdf)
        assert result == "OCR result via MarkItDown"

    def test_pdf_explicit_pymupdf4llm_skips_sparsity_check(self, tmp_path):
        """'pymupdf4llm' mode: use output as-is even if sparse."""
        pdf = tmp_path / "explicit.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        sparse_text = "x" * 10  # very short

        with (
            patch(
                "ideer.utils.file_conversion._convert_pdf_with_pymupdf4llm",
                return_value=sparse_text,
            ),
            patch("ideer.utils.file_conversion._convert_with_markitdown") as mock_md,
        ):
            result = _do_convert(pdf, "pymupdf4llm")

        mock_md.assert_not_called()
        assert result == sparse_text

    def test_pdf_explicit_markitdown_skips_pymupdf4llm(self, tmp_path):
        """'markitdown' mode: never attempt pymupdf4llm."""
        pdf = tmp_path / "force_md.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        with (
            patch("ideer.utils.file_conversion._convert_pdf_with_pymupdf4llm") as mock_pymu,
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="MarkItDown result",
            ),
        ):
            result = _do_convert(pdf, "markitdown")

        mock_pymu.assert_not_called()
        assert result == "MarkItDown result"

    def test_pdf_auto_falls_back_when_pymupdf4llm_not_installed(self, tmp_path):
        """auto mode: if pymupdf4llm is not installed, use MarkItDown directly."""
        pdf = tmp_path / "no_pymupdf.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        with (
            patch(
                "ideer.utils.file_conversion._convert_pdf_with_pymupdf4llm",
                return_value=None,  # None signals not installed
            ),
            patch(
                "ideer.utils.file_conversion._convert_with_markitdown",
                return_value="MarkItDown fallback",
            ) as mock_md,
        ):
            result = _do_convert(pdf, "auto")

        mock_md.assert_called_once_with(pdf)
        assert result == "MarkItDown fallback"


class TestGetPdfConverter:
    def test_reads_dict_backed_uploads_config(self):
        cfg = MagicMock()
        cfg.uploads = {"pdf_converter": "markitdown"}

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_pdf_converter() == "markitdown"

    def test_reads_attribute_backed_uploads_config(self):
        cfg = MagicMock()
        cfg.uploads = MagicMock(pdf_converter="pymupdf4llm")

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_pdf_converter() == "pymupdf4llm"

    def test_invalid_value_falls_back_to_auto(self):
        cfg = MagicMock()
        cfg.uploads = {"pdf_converter": "not-a-real-converter"}

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_pdf_converter() == "auto"


class TestGetDocxConverter:
    def test_reads_dict_backed_uploads_config(self):
        cfg = MagicMock()
        cfg.uploads = {"docx_converter": "markitdown"}

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_docx_converter() == "markitdown"

    def test_reads_attribute_backed_uploads_config(self):
        cfg = MagicMock()
        cfg.uploads = MagicMock(docx_converter="rich")

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_docx_converter() == "rich"

    def test_invalid_value_falls_back_to_auto(self):
        cfg = MagicMock()
        cfg.uploads = {"docx_converter": "not-a-real-converter"}

        with patch("ideer.utils.file_conversion.get_app_config", return_value=cfg):
            assert _get_docx_converter() == "auto"


class TestConvertWordWithRich:
    """Integration-style tests for the rich Word branch (real .docx fixtures)."""

    @staticmethod
    def _build_docx_with_image_and_prose(path):
        """docx with an embedded image AND prose text containing '(images/'."""
        import io
        import struct
        import zlib

        from docx import Document
        from docx.shared import Inches

        # Minimal valid 1x1 red PNG
        def chunk(tag, data):
            payload = tag + data
            return struct.pack(">I", len(data)) + payload + struct.pack(">I", zlib.crc32(payload))

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        png = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00")) + chunk(b"IEND", b"")

        doc = Document()
        doc.add_paragraph("标题一", style="Heading 1")
        para = doc.add_paragraph()
        para.add_run().add_picture(io.BytesIO(png), width=Inches(0.5))
        # Prose that literally contains "(images/" — must NOT be rewritten
        doc.add_paragraph("参见路径 (images/keep.png 获取原图)")
        doc.save(str(path))

    def test_relocates_artifacts_and_rewrites_refs(self, tmp_path):
        from ideer.utils.file_conversion import _convert_word_with_rich

        src = tmp_path / "report.docx"
        self._build_docx_with_image_and_prose(src)

        text = _convert_word_with_rich(src)

        assert text is not None
        # Image reference rewritten to the relocated directory
        assert "](report_files/images/" in text
        assert "](images/" not in text
        # Prose containing "(images/" is left untouched
        assert "(images/keep.png" in text
        # Intermediate markdown removed; artifacts live under <stem>_files/
        assert not (tmp_path / "report_files" / "report.md").exists()
        assert (tmp_path / "report_files" / "images").is_dir()

    def test_guard_rail_runtime_error_propagates(self, tmp_path):
        """Deliberate RuntimeErrors (e.g. .doc without soffice) are not masked."""
        from ideer.utils.file_conversion import _convert_word_with_rich

        src = tmp_path / "legacy.doc"
        src.write_bytes(b"\xd0\xcf\x11\xe0 fake OLE")

        with (
            patch("ideer.utils.docx_rich.is_available", return_value=True),
            patch(
                "ideer.utils.docx_rich.convert_docx",
                side_effect=RuntimeError("LibreOffice (soffice) is required"),
            ),
        ):
            with pytest.raises(RuntimeError, match="LibreOffice"):
                _convert_word_with_rich(src)


class TestConvertFileToMarkdown:
    def test_small_file_runs_synchronously(self, tmp_path):
        """Small files (< 1 MB) are converted in the event loop thread."""
        pdf = tmp_path / "small.pdf"
        pdf.write_bytes(b"%PDF-1.4 " + b"x" * 100)  # well under 1 MB

        with (
            patch("ideer.utils.file_conversion._get_pdf_converter", return_value="auto"),
            patch(
                "ideer.utils.file_conversion._do_convert",
                return_value="# Small PDF",
            ) as mock_convert,
            patch("asyncio.to_thread") as mock_thread,
        ):
            md_path = _run(convert_file_to_markdown(pdf))

        # asyncio.to_thread must NOT have been called
        mock_thread.assert_not_called()
        mock_convert.assert_called_once()
        assert md_path == pdf.with_suffix(".md")
        assert md_path.read_text() == "# Small PDF"

    def test_large_file_offloaded_to_thread(self, tmp_path):
        """Large files (> 1 MB) are offloaded via asyncio.to_thread."""
        pdf = tmp_path / "large.pdf"
        # Write slightly more than the threshold
        pdf.write_bytes(b"%PDF-1.4 " + b"x" * (_ASYNC_THRESHOLD_BYTES + 1))

        async def fake_to_thread(fn, *args, **kwargs):
            return fn(*args, **kwargs)

        with (
            patch("ideer.utils.file_conversion._get_pdf_converter", return_value="auto"),
            patch(
                "ideer.utils.file_conversion._do_convert",
                return_value="# Large PDF",
            ),
            patch("asyncio.to_thread", side_effect=fake_to_thread) as mock_thread,
        ):
            md_path = _run(convert_file_to_markdown(pdf))

        mock_thread.assert_called_once()
        assert md_path == pdf.with_suffix(".md")
        assert md_path.read_text() == "# Large PDF"

    def test_returns_none_on_conversion_error(self, tmp_path):
        """If conversion raises, return None without propagating the exception."""
        pdf = tmp_path / "broken.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        with (
            patch("ideer.utils.file_conversion._get_pdf_converter", return_value="auto"),
            patch(
                "ideer.utils.file_conversion._do_convert",
                side_effect=RuntimeError("conversion failed"),
            ),
        ):
            result = _run(convert_file_to_markdown(pdf))

        assert result is None

    def test_writes_utf8_markdown_file(self, tmp_path):
        """Generated .md file is written with UTF-8 encoding."""
        pdf = tmp_path / "report.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")
        chinese_content = "# 中文报告\n\n这是测试内容。"

        with (
            patch("ideer.utils.file_conversion._get_pdf_converter", return_value="auto"),
            patch(
                "ideer.utils.file_conversion._do_convert",
                return_value=chinese_content,
            ),
        ):
            md_path = _run(convert_file_to_markdown(pdf))

        assert md_path is not None
        assert md_path.read_text(encoding="utf-8") == chinese_content


# ---------------------------------------------------------------------------
# extract_outline
# ---------------------------------------------------------------------------


class TestExtractOutline:
    """Tests for extract_outline()."""

    def test_empty_file_returns_empty(self, tmp_path):
        """Empty markdown file yields no outline entries."""
        md = tmp_path / "empty.md"
        md.write_text("", encoding="utf-8")
        assert extract_outline(md) == []

    def test_missing_file_returns_empty(self, tmp_path):
        """Non-existent path returns [] without raising."""
        assert extract_outline(tmp_path / "nonexistent.md") == []

    def test_standard_markdown_headings(self, tmp_path):
        """# / ## / ### headings are all recognised."""
        md = tmp_path / "doc.md"
        md.write_text(
            "# Chapter One\n\nSome text.\n\n## Section 1.1\n\nMore text.\n\n### Sub 1.1.1\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        assert len(outline) == 3
        assert outline[0] == {"title": "Chapter One", "line": 1}
        assert outline[1] == {"title": "Section 1.1", "line": 5}
        assert outline[2] == {"title": "Sub 1.1.1", "line": 9}

    def test_bold_sec_item_heading(self, tmp_path):
        """**ITEM N. TITLE** lines in SEC filings are recognised."""
        md = tmp_path / "10k.md"
        md.write_text(
            "Cover page text.\n\n**ITEM 1. BUSINESS**\n\nBody.\n\n**ITEM 1A. RISK FACTORS**\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        assert len(outline) == 2
        assert outline[0] == {"title": "ITEM 1. BUSINESS", "line": 3}
        assert outline[1] == {"title": "ITEM 1A. RISK FACTORS", "line": 7}

    def test_bold_part_heading(self, tmp_path):
        """**PART I** / **PART II** headings are recognised."""
        md = tmp_path / "10k.md"
        md.write_text("**PART I**\n\n**PART II**\n\n**PART III**\n", encoding="utf-8")
        outline = extract_outline(md)
        assert len(outline) == 3
        titles = [e["title"] for e in outline]
        assert "PART I" in titles
        assert "PART II" in titles
        assert "PART III" in titles

    def test_sec_cover_page_boilerplate_excluded(self, tmp_path):
        """Address lines and short cover boilerplate must NOT appear in outline."""
        md = tmp_path / "8k.md"
        md.write_text(
            "## **UNITED STATES SECURITIES AND EXCHANGE COMMISSION**\n\n**WASHINGTON, DC 20549**\n\n**CURRENT REPORT**\n\n**SIGNATURES**\n\n**TESLA, INC.**\n\n**ITEM 2.02. RESULTS OF OPERATIONS**\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        titles = [e["title"] for e in outline]
        # Cover-page boilerplate should be excluded
        assert "WASHINGTON, DC 20549" not in titles
        assert "CURRENT REPORT" not in titles
        assert "SIGNATURES" not in titles
        assert "TESLA, INC." not in titles
        # Real SEC heading must be included
        assert "ITEM 2.02. RESULTS OF OPERATIONS" in titles

    def test_chinese_headings_via_standard_markdown(self, tmp_path):
        """Chinese annual report headings emitted as # by pymupdf4llm are captured."""
        md = tmp_path / "annual.md"
        md.write_text(
            "# 第一节 公司简介\n\n内容。\n\n## 第三节 管理层讨论与分析\n\n分析内容。\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        assert len(outline) == 2
        assert outline[0]["title"] == "第一节 公司简介"
        assert outline[1]["title"] == "第三节 管理层讨论与分析"

    def test_outline_capped_at_max_entries(self, tmp_path):
        """When truncated, result has MAX_OUTLINE_ENTRIES real entries + 1 sentinel."""
        lines = [f"# Heading {i}" for i in range(MAX_OUTLINE_ENTRIES + 10)]
        md = tmp_path / "long.md"
        md.write_text("\n".join(lines), encoding="utf-8")
        outline = extract_outline(md)
        # Last entry is the truncation sentinel
        assert outline[-1] == {"truncated": True}
        # Visible entries are exactly MAX_OUTLINE_ENTRIES
        visible = [e for e in outline if not e.get("truncated")]
        assert len(visible) == MAX_OUTLINE_ENTRIES

    def test_no_truncation_sentinel_when_under_limit(self, tmp_path):
        """Short documents produce no sentinel entry."""
        lines = [f"# Heading {i}" for i in range(5)]
        md = tmp_path / "short.md"
        md.write_text("\n".join(lines), encoding="utf-8")
        outline = extract_outline(md)
        assert len(outline) == 5
        assert not any(e.get("truncated") for e in outline)

    def test_blank_lines_and_whitespace_ignored(self, tmp_path):
        """Blank lines between headings do not produce empty entries."""
        md = tmp_path / "spaced.md"
        md.write_text("\n\n# Title One\n\n\n\n# Title Two\n\n", encoding="utf-8")
        outline = extract_outline(md)
        assert len(outline) == 2
        assert all(e["title"] for e in outline)

    def test_inline_bold_not_confused_with_heading(self, tmp_path):
        """Mid-sentence bold text must not be mistaken for a heading."""
        md = tmp_path / "prose.md"
        md.write_text(
            "This sentence has **bold words** inside it.\n\nAnother with **MULTIPLE CAPS** inline.\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        assert outline == []

    def test_split_bold_heading_academic_paper(self, tmp_path):
        """**<num>** **<title>** lines from academic papers are recognised (Style 3)."""
        md = tmp_path / "paper.md"
        md.write_text(
            "## **Attention Is All You Need**\n\n**1** **Introduction**\n\nBody text.\n\n**2** **Background**\n\nMore text.\n\n**3.1** **Encoder and Decoder Stacks**\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        titles = [e["title"] for e in outline]
        assert "1 Introduction" in titles
        assert "2 Background" in titles
        assert "3.1 Encoder and Decoder Stacks" in titles

    def test_split_bold_year_columns_excluded(self, tmp_path):
        """Financial table headers like **2023** **2022** **2021** are NOT headings."""
        md = tmp_path / "annual.md"
        md.write_text(
            "# Financial Summary\n\n**2023** **2022** **2021**\n\nRevenue 100 90 80\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        titles = [e["title"] for e in outline]
        # Only the # heading should appear, not the year-column row
        assert titles == ["Financial Summary"]

    def test_adjacent_bold_spans_merged_in_markdown_heading(self, tmp_path):
        """** ** artefacts inside a # heading are merged into clean plain text."""
        md = tmp_path / "sec.md"
        md.write_text(
            "## **UNITED STATES** **SECURITIES AND EXCHANGE COMMISSION**\n\nBody text.\n",
            encoding="utf-8",
        )
        outline = extract_outline(md)
        assert len(outline) == 1
        # Title must be clean — no ** ** artefacts
        assert outline[0]["title"] == "UNITED STATES SECURITIES AND EXCHANGE COMMISSION"
