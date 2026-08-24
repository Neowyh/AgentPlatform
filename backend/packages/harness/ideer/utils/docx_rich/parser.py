"""Rich docx parser: converts Word documents to high-fidelity Markdown.

Ported from the only_parser project. Uses python-docx to parse .docx files
directly (no LibreOffice -> PDF detour), preserving:

- Headings (Heading styles and outline levels) as Markdown headings
- Inline formatting (bold / italic / underline / monospace)
- Ordered & unordered lists
- Tables as HTML
- Embedded images extracted to an output directory
- MathType (MTEF) OLE formulas converted to LaTeX
- Visio OLE drawings rendered to PNG (requires LibreOffice)
- Legacy .doc files via LibreOffice conversion first

Formulas that MTEF cannot parse degrade to their embedded thumbnail image.
The heavy MFR (vision model) fallback from the original project is not ported.
"""

import logging
import os
import re
import shutil
from pathlib import Path

logger = logging.getLogger(__name__)


# ============================================================================
# OLE 类型注册表
# ============================================================================


class OLETypeInfo:
    """
    OLE 类型描述，定义一种 OLE 对象的识别和处理策略。

    Attributes:
        category:  类型分类（如 "formula", "visio", "chart" 等）
        label:     Markdown 输出中的图片标签名（如 "Visio绘图", "图表"）
        handler:   可选的专用提取方法名（字符串），None 表示无需特殊处理
    """

    __slots__ = ("category", "label", "handler")

    def __init__(self, category: str, label: str, handler: str | None = None):
        self.category = category
        self.label = label
        self.handler = handler


# ProgID 前缀 → OLE 类型描述
# 新增 OLE 类型只需在此表添加一行
OLE_TYPE_REGISTRY: dict[str, OLETypeInfo] = {
    "Equation": OLETypeInfo("formula", "公式"),
    "Visio.Drawing": OLETypeInfo("visio", "Visio绘图", handler="_extract_visio_ole"),
    "Visio.Stencil": OLETypeInfo("visio", "Visio绘图", handler="_extract_visio_ole"),
    # 未来扩展示例:
    # "Excel.Sheet":     OLETypeInfo("excel",  "Excel表格", handler="_extract_excel_ole"),
    # "PowerPoint.Show": OLETypeInfo("ppt",    "PPT演示",   handler="_extract_ppt_ole"),
    # "Chart":           OLETypeInfo("chart",  "图表"),
}


def _classify_prog_id(prog_id: str) -> OLETypeInfo | None:
    """
    根据 ProgID 字符串匹配 OLE 类型注册表。

    Args:
        prog_id: OLE 对象的 ProgID（如 "Visio.Drawing.11", "Equation.DSMT4"）

    Returns:
        匹配到的 OLETypeInfo，未匹配返回 None
    """
    for prefix, info in OLE_TYPE_REGISTRY.items():
        if prefix in prog_id:
            return info
    return None


class DocxDocumentParser:
    """
    基于 python-docx 的 Word 文档解析器

    将 .docx 文件直接解析为 Markdown 文本，支持：
    - 标题（Heading 1-9 映射为 # - #######）
    - 段落文本（含粗体、斜体、下划线、代码样式）
    - 有序列表和无序列表
    - 表格（HTML 格式输出，保持与现有解析器一致）
    - 图片（提取并保存到输出目录）
    """

    def ping(self) -> bool:
        """检查 python-docx 是否可用"""
        try:
            import docx  # noqa: F401  (availability probe)

            return True
        except ImportError:
            logger.warning("python-docx 未安装，无法解析 docx 文件")
            return False

    def parse_docx(self, file_path: str, output_dir: str | None = None, subdir_ext: str | None = None) -> tuple[str, dict[int, list[dict]], dict]:
        """
        解析单个 .docx 文件为 Markdown

        Args:
            file_path: .docx 文件路径
            output_dir: 输出目录（用于保存图片等资源）
            subdir_ext: 输出子目录拓展名（默认取文件自身拓展名）；
                        .doc 转换场景下传入原始拓展名，保证子目录以源文件命名

        Returns:
            (markdown_content, bbox_map, extra_info)
            - markdown_content: Markdown 文本内容
            - bbox_map: 空字典（docx 不支持 bbox）
            - extra_info: 包含图片路径等附加信息
        """
        from docx import Document as DocxDocument

        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        logger.info(f"开始解析 docx: {file_path}")

        # 读取文档
        doc = DocxDocument(file_path)

        # 准备输出子目录（以 文件名_拓展名 命名）
        file_stem = Path(file_path).stem
        ext = (subdir_ext or Path(file_path).suffix).lstrip(".").lower()
        if output_dir:
            file_output_dir = os.path.join(output_dir, f"{file_stem}_{ext}")
            os.makedirs(file_output_dir, exist_ok=True)
            image_dir = os.path.join(file_output_dir, "images")
            os.makedirs(image_dir, exist_ok=True)
        else:
            file_output_dir = None
            image_dir = None

        # 提取图片
        image_map = {}  # rId -> 图片文件路径
        latex_map = {}  # rId -> LaTeX 字符串（公式识别结果）
        if image_dir:
            image_map = self._extract_images(doc, image_dir)

            # 分类 OLE 对象（公式 vs Visio vs 其他）
            formula_rids, ole_img_rids, typed_ole_info = self._build_ole_classification(file_path)

            # 优先处理有专用 handler 的 OLE 类型（如 Visio）
            if typed_ole_info:
                self._process_typed_oles(file_path, typed_ole_info, image_dir, image_map)

            # 排除已处理的 OLE，转换剩余 OLE 缩略图为 PNG
            handled_rids = set()
            for items in typed_ole_info.values():
                handled_rids.update(items.keys())
            remaining_ole_rids = ole_img_rids - formula_rids - handled_rids
            image_map = self._convert_ole_images(image_dir, image_map, remaining_ole_rids)

            # 优先用 MTEF 二进制解析提取 LaTeX（精确，无需图像识别）
            latex_map = self._extract_mtef_latex(file_path)

            mtef_count = len(latex_map)
            logger.info(f"公式识别: MTEF={mtef_count}, 未识别(保留缩略图)={len(formula_rids) - mtef_count}")

        # 提取目录（TOC）- 从 XML 中直接提取
        toc_entries = self._extract_toc_from_xml(file_path)
        logger.info(f"从文档中提取了 {len(toc_entries)} 个目录条目")

        # 解析文档内容为 Markdown（按段落组织：文本 + 图片）
        # 使用段落索引而不是 id()，避免内存地址不匹配的问题
        para_content = {}  # {para_idx: [line1, line2, ...]}
        image_counter = [0]  # 用于生成图片引用名

        # 构建 OLE 标签映射: {img_rid: label}（用于 Markdown 输出中区分 Visio 绘图等）
        ole_label_map: dict[str, str] = {}
        for items in typed_ole_info.values():
            for rid, item in items.items():
                ole_label_map[rid] = item["info"].label

        for para_idx, para in enumerate(doc.paragraphs):
            para_lines = []

            # 段落文本（非空段落有对应的 md 行，含内联公式/OLE 缩略图）
            line = self._parse_paragraph_to_markdown(para, image_map, latex_map, file_output_dir, image_counter, ole_label_map=ole_label_map)
            if line is not None:
                para_lines.append(line)

            # 段落中的内联图片
            image_lines = self._get_paragraph_image_lines(para, image_map, file_output_dir, image_counter)
            para_lines.extend(image_lines)

            if para_lines:
                para_content[para_idx] = para_lines

        # 将表格插入到对应位置，并添加目录
        md_with_tables = self._merge_tables_into_markdown(doc, para_content, toc_entries, file_path)

        # 规范化空行：去除连续多余空行
        content = "\n".join(md_with_tables)
        content = re.sub(r"\n{3,}", "\n\n", content)
        content = content.strip()

        # 保存中间 markdown
        if file_output_dir:
            md_path = os.path.join(file_output_dir, f"{file_stem}.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content)
            logger.info(f"已保存 markdown: {md_path}")

        extra_info = {
            "image_dir": image_dir,
            "images": list(image_map.values()),
        }
        logger.info(f"docx 解析完成: {file_path}，{len(content)} 字符")

        return content, {}, extra_info

    def parse_doc(self, file_path: str, output_dir: str | None = None) -> tuple[str, dict[int, list[dict]], dict]:
        """
        解析 .doc 文件（先转换为 .docx，再解析）

        Args:
            file_path: .doc 文件路径
            output_dir: 输出目录

        Returns:
            (markdown_content, bbox_map, extra_info)
        """
        from .libreoffice import get_libreoffice_manager

        logger.info(f"将 .doc 转换为 .docx: {file_path}")

        # 创建临时目录用于存放转换后的 docx
        convert_dir = os.path.join(os.path.dirname(file_path), "_converted_docx")
        os.makedirs(convert_dir, exist_ok=True)

        libre = get_libreoffice_manager()
        docx_path = libre.convert_to_modern_format(file_path, convert_dir, timeout=120)

        if not docx_path or not os.path.exists(docx_path):
            raise RuntimeError(f".doc 转换为 .docx 失败: {file_path}")

        logger.info(f"转换成功: {docx_path}")

        try:
            # 再用 docx 解析器解析（子目录以原始 .doc 拓展名命名）
            return self.parse_docx(docx_path, output_dir, subdir_ext=Path(file_path).suffix)
        finally:
            # 清理临时转换目录，避免在上传目录中残留垃圾
            shutil.rmtree(convert_dir, ignore_errors=True)

    def parse_docx_batch(self, docx_files: list[str], output_dir: str) -> dict[str, dict]:
        """
        批量解析多个 .docx 文件

        Args:
            docx_files: .docx 文件路径列表
            output_dir: 输出目录

        Returns:
            结果字典: {文件名: {"file": str, "success": bool, "error": Optional[str], "markdown_path": Optional[str]}}
        """
        if not self.ping():
            return {os.path.basename(fp): {"file": fp, "success": False, "error": "python-docx 不可用"} for fp in docx_files}

        results = {}
        for idx, fp in enumerate(docx_files, 1):
            logger.info(f"[{idx}/{len(docx_files)}] 解析: {os.path.basename(fp)}")
            try:
                markdown_content, bbox_map, extra_info = self.parse_docx(fp, output_dir)
                results[os.path.basename(fp)] = {
                    "file": fp,
                    "success": True,
                }
                logger.info(f"  ✓ {os.path.basename(fp)}")
            except Exception as e:
                logger.error(f"  ✗ {os.path.basename(fp)}: {e}")
                results[os.path.basename(fp)] = {
                    "file": fp,
                    "success": False,
                    "error": str(e),
                }

        success = sum(1 for r in results.values() if r.get("success"))
        logger.info(f"docx 批量解析完成: 成功 {success}/{len(docx_files)}")
        return results

    # ========================================================================
    # 内部方法
    # ========================================================================

    def _extract_toc_from_xml(self, docx_path: str) -> list[dict]:
        """
        从 DOCX 的 XML 中提取目录（TOC）

        处理超链接类型的目录条目

        Args:
            docx_path: DOCX 文件路径

        Returns:
            [{"text": "目录文本", "level": 1}, ...] 列表
        """
        import zipfile

        from lxml import etree

        toc_data = []

        try:
            with zipfile.ZipFile(docx_path) as z:
                with z.open("word/document.xml") as f:
                    tree = etree.parse(f)
                    root = tree.getroot()

                    # 定义 Word 的 XML 命名空间
                    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main", "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}

                    # 查找所有的结构化标签内容 (SDT Content)，也就是目录的容器
                    sdt_contents = root.findall(".//w:sdtContent", namespaces)

                    for sdt in sdt_contents:
                        # 在目录容器内查找所有的段落
                        paragraphs = sdt.findall(".//w:p", namespaces)

                        for p in paragraphs:
                            # 提取纯文本（包括超链接内的文本）
                            # 使用 xpath 提取所有 w:t 标签的文本
                            texts = p.xpath(".//w:t/text()", namespaces=namespaces)
                            line_text = "".join(texts).strip()

                            if not line_text:
                                continue

                            # 提取层级样式 (例如 w:val="TOC 1")
                            p_style_elem = p.find(".//w:pStyle", namespaces)
                            level = 0
                            if p_style_elem is not None:
                                style_val = p_style_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "")
                                if style_val.startswith("TOC"):
                                    # 提取 TOC 1, TOC 2 等后面的数字
                                    level_str = style_val.split()[-1] if style_val.split() else "1"
                                    level = int(level_str) if level_str.isdigit() else 1

                            toc_data.append({"text": line_text, "level": level})

        except Exception as e:
            logger.warning(f"提取目录失败: {e}")

        logger.info(f"成功提取 {len(toc_data)} 个目录条目: {toc_data[:3] if toc_data else 'none'}...")
        return toc_data

    def _parse_paragraph_to_markdown(
        self, para, image_map: dict[str, str] | None = None, latex_map: dict[str, str] | None = None, output_base_dir: str | None = None, counter: list[int] | None = None, ole_label_map: dict[str, str] | None = None
    ) -> str | None:
        """
        将 python-docx 段落转换为 Markdown 行

        Args:
            para: python-docx Paragraph 对象
            image_map: {rId: 图片文件绝对路径} 映射（用于 OLE 缩略图）
            latex_map: {rId: LaTeX 字符串} 映射（公式识别结果，优先使用）
            output_base_dir: markdown 文件所在目录（用于计算相对路径）
            counter: 图片计数器
            ole_label_map: {img_rid: label} OLE 类型标签映射（如 "Visio绘图"）

        Returns:
            Markdown 字符串，空段落返回 None
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        text = para.text.strip()

        # 检查段落中是否有 OLE 对象（即使没有文本，有 OLE 对象也要处理）
        has_ole = False
        if not text:
            ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for run in para.runs:
                if run._element.findall(f"{{{ns_w}}}object"):
                    has_ole = True
                    break
            if not has_ole:
                return None

        # 安全获取样式名
        try:
            style_name = para.style.name.lower() if para.style else ""
        except Exception:
            style_name = ""

        # 检测标题级别（标准 Heading 样式）
        heading_match = re.match(r"heading\s*(\d+)", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            level = min(level, 6)  # Markdown 最多 6 级标题
            return f"{'#' * level} {text}"

        # 检测自动编号的标题（通过大纲级别）
        # 注意：大纲级别 0-8 是标题，9 是正文
        try:
            pPr = para._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
            if pPr is not None:
                outlineLvl = pPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl")
                if outlineLvl is not None:
                    try:
                        val = outlineLvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                        if val is not None:
                            level = int(val)
                            # 只有 0-8 才是标题级别，9 是正文
                            if level < 9:
                                level = level + 1  # 转换为 Markdown 级别（1-based）
                                level = min(level, 6)
                                return f"{'#' * level} {text}"
                    except (ValueError, AttributeError, TypeError):
                        pass
        except Exception:
            pass

        # 检测列表（通过 numPr 属性）
        try:
            numPr = para._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if numPr is not None:
                # 有序或无序列表
                # 检测编号格式来判断有序/无序
                numFmt = numPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt")
                if numFmt is not None and numFmt.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") in ("decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"):
                    ordered_index = self._get_list_index(para)
                    return f"  {ordered_index}. {text}"
                else:
                    return f"- {text}"
        except Exception:
            pass

        # 检测段落样式中的列表（ListBullet, ListNumber 等）
        if "list" in style_name:
            return f"- {text}"

        # 安全获取对齐方式
        try:
            alignment = para.alignment
            is_center = alignment == WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            is_center = False

        if is_center:
            return text

        # 默认情况：处理为普通段落文本（保留内联格式，含公式 LaTeX / OLE 缩略图）
        try:
            formatted_text = self._get_formatted_paragraph_text(para, image_map, latex_map, output_base_dir, counter, ole_label_map=ole_label_map)
            # 确保返回非空内容
            if formatted_text:
                return formatted_text
        except Exception:
            pass

        # 降级：直接返回纯文本
        return text

    def _get_formatted_paragraph_text(
        self, para, image_map: dict[str, str] | None = None, latex_map: dict[str, str] | None = None, output_base_dir: str | None = None, counter: list[int] | None = None, ole_label_map: dict[str, str] | None = None
    ) -> str:
        """
        提取段落的格式化文本（保留粗体、斜体、代码等），并内联插入公式 LaTeX 或 OLE 缩略图

        Args:
            para: python-docx Paragraph 对象
            image_map: {rId: 图片文件绝对路径} 映射（用于 OLE 缩略图）
            latex_map: {rId: LaTeX 字符串} 映射（公式识别结果，优先使用）
            output_base_dir: markdown 文件所在目录（用于计算相对路径）
            counter: 图片计数器
            ole_label_map: {img_rid: label} OLE 类型标签映射（如 "Visio绘图"）

        Returns:
            带 Markdown 格式标记的文本（含内联 LaTeX 或图片）
        """
        parts = []
        for run in para.runs:
            text = run.text
            if text:
                # 安全获取字体样式属性（部分 docx 中枚举值可能为空或不合法）
                bold = self._safe_get_run_attr(run, "bold")
                italic = self._safe_get_run_attr(run, "italic")
                self._safe_get_run_attr(run, "underline")
                font_name = self._safe_get_font_name(run)

                # 代码样式：等宽字体或显式代码样式
                is_code = False
                if font_name and font_name.lower() in ("consolas", "courier new", "monospace", "courier"):
                    is_code = True

                # 应用 Markdown 格式
                formatted = text
                if is_code:
                    formatted = f"`{text}`"
                if bold:
                    formatted = f"**{formatted}**"
                if italic:
                    formatted = f"*{formatted}*"

                parts.append(formatted)

            # 处理 OLE 对象（如 MathType 公式）：优先 LaTeX，降级为缩略图
            if (image_map or latex_map) and counter:
                ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                ns_v = "urn:schemas-microsoft-com:vml"
                ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

                w_objects = run._element.findall(f"{{{ns_w}}}object")
                for w_obj in w_objects:
                    imagedata_list = w_obj.findall(f".//{{{ns_v}}}imagedata")
                    for imagedata in imagedata_list:
                        img_rid = imagedata.get(f"{{{ns_r}}}id")

                        # 优先使用 LaTeX（公式识别成功）
                        if img_rid and latex_map and img_rid in latex_map:
                            # 占位，后面根据段落是否纯公式决定行内/行间
                            parts.append(("LATEX", latex_map[img_rid]))
                        # 降级：使用缩略图
                        elif img_rid and image_map and img_rid in image_map:
                            counter[0] += 1
                            abs_path = image_map[img_rid]
                            if output_base_dir:
                                rel_path = os.path.relpath(abs_path, output_base_dir).replace("\\", "/")
                            else:
                                rel_path = abs_path.replace("\\", "/")
                            # 使用注册表标签区分 OLE 类型（如 "Visio绘图"）
                            label = (ole_label_map or {}).get(img_rid, "image")
                            parts.append(f"![{label}-{counter[0]}]({rel_path})")

        # 判断是否为纯公式段落（无文本内容，只有公式）
        has_text = any(not isinstance(p, tuple) and p.strip() for p in parts)

        # 将 LaTeX 占位替换为实际的行内/行间格式
        result_parts = []
        for p in parts:
            if isinstance(p, tuple) and p[0] == "LATEX":
                if has_text:
                    result_parts.append(f"${p[1]}$")
                else:
                    result_parts.append(f"$${p[1]}$$")
            else:
                result_parts.append(p)

        return "".join(result_parts) if result_parts else para.text

    @staticmethod
    def _safe_get_run_attr(run, attr: str):
        """
        安全获取 run 的属性，避免枚举映射异常

        Args:
            run: python-docx Run 对象
            attr: 属性名（bold, italic, underline 等）

        Returns:
            属性值，异常时返回 None
        """
        try:
            return getattr(run, attr, None)
        except Exception:
            return None

    @staticmethod
    def _safe_get_font_name(run) -> str | None:
        """
        安全获取 run 的字体名称

        Args:
            run: python-docx Run 对象

        Returns:
            字体名称，异常时返回 None
        """
        try:
            if run.font:
                return run.font.name
        except Exception:
            pass
        return None

    def _get_list_index(self, para) -> int:
        """
        获取有序列表项序号

        Args:
            para: python-docx Paragraph 对象

        Returns:
            列表序号
        """
        try:
            numPr = para._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if numPr is not None:
                ilvl = numPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
                numId = numPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
                if ilvl is not None and numId is not None:
                    # 由于完整解析编号较复杂，使用简单策略
                    return 1
        except Exception:
            pass
        return 1

    def _get_paragraph_image_lines(self, para, image_map: dict[str, str], output_base_dir: str | None, counter: list[int]) -> list[str]:
        """
        获取段落中内联图片的 Markdown 行

        Args:
            para: python-docx Paragraph 对象
            image_map: {rId: 图片文件绝对路径} 映射
            output_base_dir: markdown 文件所在目录（用于计算相对路径）
            counter: 图片计数器

        Returns:
            图片 Markdown 行列表（每张图片占两行：![...](...) + 空行）
        """
        lines = []
        for run in para.runs:
            # --- 处理普通图片（w:drawing） ---
            drawings = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
            if not drawings:
                drawings = run._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor")
            if not drawings:
                drawings = run._element.findall(".//" + self._ns_tag("w:drawing"))

            for _ in drawings:
                # 查找图片引用
                blip = run._element.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                if blip is not None:
                    embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed and embed in image_map:
                        counter[0] += 1
                        abs_path = image_map[embed]
                        # 计算相对于 markdown 文件所在目录的路径
                        if output_base_dir:
                            rel_path = os.path.relpath(abs_path, output_base_dir).replace("\\", "/")
                        else:
                            rel_path = abs_path.replace("\\", "/")
                        lines.append(f"![image-{counter[0]}]({rel_path})")
                        lines.append("")

        return lines

    @staticmethod
    def _ns_tag(tag: str) -> str:
        """
        构造带命名空间的 XML 标签名

        Args:
            tag: 不带命名空间的标签名

        Returns:
            带命名空间的标签名
        """
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        return f"{ns}{tag.split(':')[-1]}" if ":" in tag else f"{ns}{tag}"

    def _extract_images(self, doc, image_dir: str) -> dict[str, str]:
        """
        从 docx 中提取所有图片

        Args:
            doc: python-docx Document 对象
            image_dir: 图片输出目录

        Returns:
            {rId: 图片文件路径} 映射
        """
        image_map = {}

        # 遍历文档的所有关系，找出图片
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                rId = rel.rId
                image_data = rel.target_part.blob
                # 获取图片扩展名
                content_type = rel.target_part.content_type
                ext = self._content_type_to_ext(content_type)
                img_filename = f"{rId}{ext}"
                img_path = os.path.join(image_dir, img_filename)

                # 保存图片
                with open(img_path, "wb") as f:
                    f.write(image_data)

                image_map[rId] = img_path
                logger.debug(f"提取图片: {img_path} ({len(image_data)} 字节)")

        if image_map:
            logger.info(f"从 docx 中提取了 {len(image_map)} 张图片")

        return image_map

    def _extract_mtef_latex(self, docx_path: str) -> dict[str, str]:
        """
        优先通过 MTEF 二进制解析提取 MathType OLE 公式的 LaTeX

        MTEF 直接从 OLE 对象内部读取结构化公式数据，100% 精确。
        仅对 ProgID 包含 Equation 的 OLE 对象生效。

        Args:
            docx_path: DOCX 文件路径

        Returns:
            {img_rId: latex_string} 映射
        """
        try:
            from .mtef_parser import extract_latex_from_docx

            mtef_map = extract_latex_from_docx(docx_path)
            if mtef_map:
                logger.info(f"MTEF 二进制解析: {len(mtef_map)} 个公式成功转为 LaTeX")
            return mtef_map
        except ImportError:
            logger.debug("mtef_parser 模块不可用，跳过 MTEF 解析")
            return {}
        except Exception as e:
            logger.warning(f"MTEF 二进制解析失败，公式将保留为缩略图: {e}")
            return {}

    def _build_ole_classification(self, docx_path: str) -> tuple[set, set, dict[str, dict]]:
        """
        解析 document.xml，基于 OLE_TYPE_REGISTRY 分类 OLE 对象

        在 w:object 元素内：
        - o:OLEObject 的 ProgID 标识 OLE 类型
        - 同级的 v:imagedata r:id 引用缩略图

        Args:
            docx_path: DOCX 文件路径

        Returns:
            (formula_rids, ole_img_rids, typed_ole_info)
            - formula_rids: 公式类缩略图 rId 集合
            - ole_img_rids: 所有 OLE 缩略图 rId 集合
            - typed_ole_info: {category: {img_rid: {"ole_rid": str, "prog_id": str, "info": OLETypeInfo}}}
              有专用 handler 的非公式 OLE 对象，按 category 分组
        """
        import zipfile

        from lxml import etree

        formula_rids = set()
        ole_img_rids = set()
        typed_ole_info: dict[str, dict] = {}  # {category: {img_rid: {...}}}

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_v = "urn:schemas-microsoft-com:vml"
        ns_o = "urn:schemas-microsoft-com:office:office"
        ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        try:
            with zipfile.ZipFile(docx_path) as z:
                with z.open("word/document.xml") as f:
                    tree = etree.parse(f)
                    root = tree.getroot()

                    for w_obj in root.findall(f".//{{{ns_w}}}object"):
                        # 获取缩略图 rId
                        img_rid = None
                        imagedata = w_obj.find(f".//{{{ns_v}}}imagedata")
                        if imagedata is not None:
                            img_rid = imagedata.get(f"{{{ns_r}}}id")

                        if not img_rid:
                            continue

                        ole_img_rids.add(img_rid)

                        # 获取 ProgID 和 OLE 数据 rId
                        ole_elem = w_obj.find(f".//{{{ns_o}}}OLEObject")
                        if ole_elem is not None:
                            prog_id = ole_elem.get("ProgID", "")
                            ole_rid = ole_elem.get(f"{{{ns_r}}}id", "")

                            if not prog_id:
                                continue

                            # 查注册表分类
                            type_info = _classify_prog_id(prog_id)
                            if type_info is None:
                                continue

                            if type_info.category == "formula":
                                formula_rids.add(img_rid)
                            elif type_info.handler:
                                # 有专用 handler 的类型，收集信息
                                category = type_info.category
                                if category not in typed_ole_info:
                                    typed_ole_info[category] = {}
                                typed_ole_info[category][img_rid] = {
                                    "ole_rid": ole_rid,
                                    "prog_id": prog_id,
                                    "info": type_info,
                                }

        except Exception as e:
            logger.warning(f"解析 OLE 分类失败: {e}")

        # 日志
        parts = [f"{len(formula_rids)} 个公式"]
        for cat, items in typed_ole_info.items():
            label = next((i["info"].label for i in items.values()), cat)
            parts.append(f"{len(items)} 个{label}")
        other_count = len(ole_img_rids - formula_rids - set(rid for items in typed_ole_info.values() for rid in items))
        if other_count:
            parts.append(f"{other_count} 个其他")
        logger.info(f"OLE 分类: {', '.join(parts)}")

        return formula_rids, ole_img_rids, typed_ole_info

    # ------------------------------------------------------------------
    # 通用 OLE 类型分发处理
    # ------------------------------------------------------------------

    def _process_typed_oles(self, docx_path: str, typed_ole_info: dict[str, dict], image_dir: str, image_map: dict[str, str]) -> dict[str, str]:
        """
        根据 OLE 类型注册表分发到专用 handler 进行处理。

        遍历 typed_ole_info 中每个 category，调用对应 handler 方法
        （handler 方法名由 OLETypeInfo.handler 指定）。

        Args:
            docx_path: DOCX 文件路径
            typed_ole_info: {category: {img_rid: {"ole_rid": str, "prog_id": str, "info": OLETypeInfo}}}
            image_dir: 图片输出目录
            image_map: {rId: 图片文件路径} 映射（handler 可更新此映射）

        Returns:
            更新后的 image_map
        """
        handled_rids = set()

        for category, items in typed_ole_info.items():
            if not items:
                continue

            # 获取 handler 方法名（同一 category 内所有条目共享同一个 handler）
            sample_info = next(iter(items.values()))
            handler_name = sample_info["info"].handler
            label = sample_info["info"].label

            if not handler_name:
                continue

            handler = getattr(self, handler_name, None)
            if handler is None:
                logger.warning(f"OLE handler '{handler_name}' 不存在，跳过 {len(items)} 个{label}")
                continue

            logger.info(f"[OLE] 处理 {len(items)} 个{label} (handler={handler_name})")
            try:
                updated_rids = handler(docx_path, items, image_dir, image_map)
                if updated_rids:
                    handled_rids.update(updated_rids)
            except Exception as e:
                logger.error(f"[OLE] {label}处理失败: {e}")

        return image_map

    # ------------------------------------------------------------------
    # Visio OLE 提取
    # ------------------------------------------------------------------

    def _extract_visio_ole(self, docx_path: str, visio_items: dict[str, dict], image_dir: str, image_map: dict[str, str]) -> set:
        """
        从 Visio OLE 对象中提取高清图片

        策略（按优先级）:
        1. 用 olefile 从 OLE .bin 提取 previewEnhMeta 流（EMF 预览）
        2. 降级：用 olefile 从 \\x01Ole10Native 提取原始 .vsd/.vsdx
        3. 降级：使用已有的 WMF/EMF 缩略图
        提取后用 LibreOffice 转 PNG。

        Args:
            docx_path: DOCX 文件路径
            visio_items: {img_rid: {"ole_rid": str, "prog_id": str, "info": OLETypeInfo}}
            image_dir: 图片输出目录
            image_map: {rId: 图片文件路径} 映射（会被更新）

        Returns:
            成功处理的 img_rid 集合
        """
        import zipfile

        handled_rids = set()
        files_to_convert = []  # [(img_rid, file_path)]

        # 解析 docx 中的关系映射，获取 OLE .bin 文件路径
        ole_path_map = self._resolve_ole_paths(docx_path, visio_items)

        try:
            with zipfile.ZipFile(docx_path) as z:
                for img_rid, item in visio_items.items():
                    ole_rid = item["ole_rid"]
                    prog_id = item["prog_id"]

                    # 获取 OLE .bin 文件在 ZIP 中的路径
                    ole_bin_path = ole_path_map.get(ole_rid)
                    if not ole_bin_path:
                        logger.debug(f"  [Visio] {img_rid}: 未找到 OLE 数据路径，保留缩略图")
                        continue

                    # 读取 OLE .bin 数据
                    try:
                        ole_data = z.read(ole_bin_path)
                    except KeyError:
                        logger.debug(f"  [Visio] {img_rid}: ZIP 中不存在 {ole_bin_path}")
                        continue

                    extracted_file = self._extract_visio_from_ole(ole_data, img_rid, prog_id, image_dir)

                    if extracted_file:
                        files_to_convert.append((img_rid, extracted_file))
                        logger.info(f"  [Visio] {img_rid}: 提取成功 ({os.path.basename(extracted_file)})")

        except Exception as e:
            logger.warning(f"[Visio] ZIP 读取失败: {e}")

        # 用 LibreOffice 批量转换为 PNG
        if files_to_convert:
            self._convert_visio_to_png(files_to_convert, image_dir, image_map)
            handled_rids = {rid for rid, _ in files_to_convert}

        return handled_rids

    @staticmethod
    def _resolve_ole_paths(docx_path: str, items: dict[str, dict]) -> dict[str, str]:
        """
        解析 OLE 数据文件在 ZIP 中的路径

        从 word/_rels/document.xml.rels 中查找 ole_rid → embeddings/xxx.bin 的映射。

        Args:
            docx_path: DOCX 文件路径
            items: {img_rid: {"ole_rid": str, ...}}

        Returns:
            {ole_rid: "word/embeddings/xxx.bin"} 映射
        """
        import zipfile

        from lxml import etree

        result = {}

        ole_rids = {item["ole_rid"] for item in items.values() if item["ole_rid"]}

        if not ole_rids:
            return result

        try:
            with zipfile.ZipFile(docx_path) as z:
                with z.open("word/_rels/document.xml.rels") as f:
                    tree = etree.parse(f)
                    root = tree.getroot()

                    for rel in root:
                        rid = rel.get("Id", "")
                        target = rel.get("Target", "")
                        if rid in ole_rids and target:
                            # Target 通常是 "embeddings/oleObject1.bin"，需要加上 "word/" 前缀
                            if not target.startswith("word/"):
                                target = "word/" + target
                            result[rid] = target
        except Exception as e:
            logger.warning(f"解析 OLE 关系映射失败: {e}")

        return result

    @staticmethod
    def _extract_visio_from_ole(ole_data: bytes, img_rid: str, prog_id: str, image_dir: str) -> str | None:
        """
        从 OLE 二进制数据中提取 Visio 预览或原始文件

        优先级:
        1. previewEnhMeta 流 → EMF 预览
        2. \\x01Ole10Native 流 → 原始 .vsd/.vsdx

        Args:
            ole_data: OLE .bin 文件的原始字节
            img_rid: 缩略图 rId（用于命名输出文件）
            prog_id: OLE 对象的 ProgID
            image_dir: 输出目录

        Returns:
            提取出的文件路径，失败返回 None
        """
        import io

        try:
            import olefile
        except ImportError:
            logger.debug("[Visio] olefile 未安装，无法提取 OLE 数据")
            return None

        try:
            ole = olefile.OleFileIO(io.BytesIO(ole_data))
        except Exception as e:
            logger.debug(f"  [Visio] {img_rid}: OLE 文件解析失败 ({e})")
            return None

        output_path = None

        try:
            # 策略 1: 提取 EMF 预览
            if ole.exists("previewEnhMeta"):
                emf_data = ole.openstream("previewEnhMeta").read()
                if emf_data and len(emf_data) > 100:
                    emf_path = os.path.join(image_dir, f"visio_{img_rid}.emf")
                    with open(emf_path, "wb") as f:
                        f.write(emf_data)
                    output_path = emf_path
                    logger.debug(f"  [Visio] {img_rid}: 提取 previewEnhMeta ({len(emf_data)} bytes)")

            # 策略 2: 提取原始 Visio 文件
            if output_path is None and ole.exists("\x01Ole10Native"):
                native_data = ole.openstream("\x01Ole10Native").read()
                if native_data:
                    vsd_sig = b"Visio (TM) Drawing"
                    zip_sig = b"PK\x03\x04"

                    if vsd_sig in native_data[:200]:
                        vsd_path = os.path.join(image_dir, f"visio_{img_rid}.vsd")
                        with open(vsd_path, "wb") as f:
                            f.write(native_data)
                        output_path = vsd_path
                        logger.debug(f"  [Visio] {img_rid}: 提取 VSD 原始文件 ({len(native_data)} bytes)")
                    else:
                        zip_offset = native_data.find(zip_sig)
                        if zip_offset >= 0:
                            vsdx_data = native_data[zip_offset:]
                            vsdx_path = os.path.join(image_dir, f"visio_{img_rid}.vsdx")
                            with open(vsdx_path, "wb") as f:
                                f.write(vsdx_data)
                            output_path = vsdx_path
                            logger.debug(f"  [Visio] {img_rid}: 提取 VSDX 原始文件 ({len(vsdx_data)} bytes)")
        except Exception as e:
            logger.debug(f"  [Visio] {img_rid}: OLE 流提取失败 ({e})")
        finally:
            ole.close()

        return output_path

    def _convert_visio_to_png(self, files_to_convert: list[tuple[str, str]], image_dir: str, image_map: dict[str, str]):
        """
        将提取出的 Visio 文件（EMF/VSD/VSDX）批量转为 PNG

        Args:
            files_to_convert: [(img_rid, extracted_file_path)]
            image_dir: 图片输出目录
            image_map: {rId: 图片文件路径}（会被更新）
        """
        try:
            from .libreoffice import get_libreoffice_manager

            libre = get_libreoffice_manager()
            soffice_path = libre._find_soffice()
            if not soffice_path:
                logger.warning("[Visio] LibreOffice 不可用，无法转换为 PNG")
                return
        except Exception:
            logger.warning("[Visio] LibreOffice 不可用，无法转换为 PNG")
            return

        files = [fp for _, fp in files_to_convert]
        self._batch_convert_with_libreoffice(soffice_path, files, image_dir, image_map, "png")

        # 裁剪空白并更新 image_map
        for img_rid, src_path in files_to_convert:
            png_path = os.path.splitext(src_path)[0] + ".png"
            if os.path.exists(png_path):
                image_map[img_rid] = png_path
                if src_path != png_path:
                    try:
                        os.remove(src_path)
                    except OSError:
                        pass

        # 裁剪空白
        png_files = [os.path.splitext(fp)[0] + ".png" for _, fp in files_to_convert]
        existing_pngs = [p for p in png_files if os.path.exists(p)]
        if existing_pngs:
            self._crop_formula_whitespace(existing_pngs, image_map)

        success = sum(1 for _, fp in files_to_convert if os.path.exists(os.path.splitext(fp)[0] + ".png"))
        logger.info(f"[Visio] LibreOffice 转换: {success}/{len(files_to_convert)} 成功")

    def _convert_ole_images(self, image_dir: str, image_map: dict[str, str], ole_img_rids: set) -> dict[str, str]:
        """
        将 OLE 缩略图统一转换为 PNG（使用 LibreOffice）

        Args:
            image_dir: 图片输出目录
            image_map: {rId: 图片文件路径} 映射
            ole_img_rids: 所有 OLE 缩略图 rId 集合

        Returns:
            更新后的 image_map
        """
        # 收集需要转换的 WMF/EMF 文件
        ole_files = [image_map[rid] for rid in ole_img_rids if rid in image_map and image_map[rid].lower().endswith((".wmf", ".emf"))]
        if not ole_files:
            return image_map

        try:
            from .libreoffice import get_libreoffice_manager

            libre = get_libreoffice_manager()
            soffice_path = libre._find_soffice()
            if soffice_path:
                self._batch_convert_with_libreoffice(soffice_path, ole_files, image_dir, image_map, "png")

                # 裁剪 PNG 空白（LibreOffice 按 A4 页面渲染，公式只占中间一小块）
                self._crop_formula_whitespace(ole_files, image_map)
        except Exception as e:
            logger.warning(f"OLE 图片格式转换失败: {e}")

        return image_map

    @staticmethod
    def _crop_formula_whitespace(ole_files: list[str], image_map: dict[str, str]):
        """
        裁剪 PNG 公式图片周围的空白

        LibreOffice 将 WMF 按 A4 页面渲染，公式只占中间一小块，
        大片空白会干扰 MFR 模型识别。使用 numpy 向量化查找非白色区域。
        """
        import numpy as np
        from PIL import Image

        cropped_count = 0
        padding = 10
        threshold = 250

        for src_path in ole_files:
            png_path = os.path.splitext(src_path)[0] + ".png"
            if not os.path.exists(png_path):
                continue
            try:
                img = Image.open(png_path)
                img_np = np.array(img)
                # 查找非白色像素（任一通道 < threshold）
                non_white = np.where(np.any(img_np < threshold, axis=-1))
                if non_white[0].size > 0:
                    y0 = max(0, non_white[0].min() - padding)
                    y1 = min(img.height, non_white[0].max() + padding + 1)
                    x0 = max(0, non_white[1].min() - padding)
                    x1 = min(img.width, non_white[1].max() + padding + 1)
                    if (img.width - (x1 - x0)) > 20 or (img.height - (y1 - y0)) > 20:
                        img.crop((x0, y0, x1, y1)).save(png_path)
                        cropped_count += 1
            except Exception:
                pass

        if cropped_count > 0:
            logger.info(f"裁剪公式图片空白: {cropped_count} 张")

    def _batch_convert_with_libreoffice(self, soffice_path: str, files: list[str], image_dir: str, image_map: dict[str, str], target_format: str) -> int:
        """
        使用 LibreOffice 批量转换图片格式（降级方案）

        Args:
            soffice_path: LibreOffice 可执行文件路径
            files: 待转换文件路径列表
            image_dir: 输出目录
            image_map: 要更新的 {rId: 路径} 映射
            target_format: 目标格式（png）

        Returns:
            成功转换的文件数
        """
        import subprocess

        batch_size = 50
        converted_count = 0

        for i in range(0, len(files), batch_size):
            batch = files[i : i + batch_size]

            cmd = [
                soffice_path,
                "--headless",
                "--convert-to",
                target_format,
                "--outdir",
                image_dir,
            ] + batch

            logger.debug(f"批量转换 WMF/EMF -> {target_format.upper()}: 第 {i // batch_size + 1} 批，{len(batch)} 个文件")

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

            if result.returncode != 0:
                logger.warning(f"格式转换部分失败: {result.stderr}")

            for src_path in batch:
                dst_path = os.path.splitext(src_path)[0] + f".{target_format}"
                if os.path.exists(dst_path):
                    for rId, path in list(image_map.items()):
                        if path == src_path:
                            image_map[rId] = dst_path
                    try:
                        os.remove(src_path)
                    except OSError:
                        pass
                    converted_count += 1

        return converted_count

    @staticmethod
    def _content_type_to_ext(content_type: str) -> str:
        """
        根据 MIME 类型返回图片扩展名

        Args:
            content_type: MIME 类型字符串

        Returns:
            文件扩展名（含.）
        """
        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/x-emf": ".emf",
            "image/x-wmf": ".wmf",
            "image/svg+xml": ".svg",
        }
        return ext_map.get(content_type, ".png")

    def _parse_table_to_markdown(self, table) -> str:
        """
        将 python-docx 表格转换为 HTML 字符串

        Args:
            table: python-docx Table 对象

        Returns:
            HTML 字符串
        """
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)
            rows.append(cells)

        if not rows:
            return ""

        # 构建 HTML 表格
        html = "<table>\n"
        for row_idx, row_cells in enumerate(rows):
            html += "  <tr>\n"
            for cell_text in row_cells:
                tag = "th" if row_idx == 0 else "td"
                html += f"    <{tag}>{cell_text}</{tag}>\n"
            html += "  </tr>\n"
        html += "</table>"

        return html

    def _merge_tables_into_markdown(self, doc, para_content: dict[int, list[str]], toc_entries: list[dict] = None, docx_path: str = None) -> list[str]:
        """
        将文档中的表格合并到 Markdown 内容中，并在 SDT 位置插入目录

        通过按 body 子元素顺序遍历（段落+表格），
        用段落索引匹配 para_content，精确定位每个段落的输出。

        Args:
            doc: python-docx Document 对象
            para_content: {para_idx: [line1, line2, ...]}
                         每个非空段落对应的 markdown 行列表（含文本和图片）
            toc_entries: 目录条目列表 [{"text": "...", "level": 1}, ...]
            docx_path: DOCX 文件路径（用于定位 SDT 位置）

        Returns:
            合并后的 Markdown 行列表
        """
        result = []

        # 标记是否已插入目录
        toc_inserted = False
        sdt_found = False  # 调试用

        # 遍历 body 的所有子元素
        body_children = list(doc.element.body)
        para_idx = 0  # 段落索引计数器

        for child in body_children:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            # 调试：记录遇到的标签
            if tag == "sdt":
                sdt_found = True
                logger.debug(f"遇到 SDT 标签，toc_entries={len(toc_entries) if toc_entries else 0}, toc_inserted={toc_inserted}")

            # 检查是否是 SDT 标签（目录容器）
            if tag == "sdt" and toc_entries and not toc_inserted:
                result.append("# 目录")
                result.append("")
                for entry in toc_entries:
                    text = entry.get("text", "").strip()
                    level = entry.get("level", 1)

                    # 根据级别生成缩进
                    indent = "  " * (level - 1)
                    result.append(f"{indent}- {text}")

                result.append("")  # 目录后空行
                toc_inserted = True
                continue

            if tag == "p":
                # 段落节点 -> 通过索引查找对应的内容行
                if para_idx in para_content:
                    result.extend(para_content[para_idx])
                # 无论是否匹配，段落索引都要递增
                para_idx += 1

            elif tag == "tbl":
                # 表格节点 -> 转换为 HTML
                table = None
                for tbl in doc.tables:
                    if tbl._element is child:
                        table = tbl
                        break

                if table is not None:
                    table_html = self._parse_table_to_markdown(table)
                    if table_html:
                        result.append(table_html)
                        result.append("")

        logger.info(f"合并完成: 共 {len(result)} 行，找到 SDT={sdt_found}, 插入目录={toc_inserted}")

        return result


# ============================================================================
# 便捷函数
# ============================================================================


def parse_docx_file(file_path: str, output_dir: str | None = None) -> tuple[str, dict[int, list[dict]], dict]:
    """
    解析 .docx 或 .doc 文件（便捷函数）

    Args:
        file_path: Word 文档路径
        output_dir: 输出目录（可选）

    Returns:
        (markdown_content, bbox_map, extra_info)
    """
    ext = Path(file_path).suffix.lower()
    parser = DocxDocumentParser()

    if ext == ".docx":
        return parser.parse_docx(file_path, output_dir)
    elif ext == ".doc":
        return parser.parse_doc(file_path, output_dir)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def batch_parse_docx_files(file_paths: list[str], output_dir: str) -> dict[str, dict]:
    """
    批量解析 Word 文档文件（.docx 和 .doc）

    Args:
        file_paths: Word 文档路径列表
        output_dir: 输出目录

    Returns:
        结果字典: {文件名: {"file": str, "success": bool, "error": Optional[str]}}
    """
    # 分离 docx 和 doc
    docx_files = [fp for fp in file_paths if Path(fp).suffix.lower() == ".docx"]
    doc_files = [fp for fp in file_paths if Path(fp).suffix.lower() == ".doc"]

    parser = DocxDocumentParser()
    results = {}

    # 处理 docx 文件
    if docx_files:
        logger.info(f"解析 {len(docx_files)} 个 .docx 文件...")
        docx_results = parser.parse_docx_batch(docx_files, output_dir)
        results.update(docx_results)

    # 处理 doc 文件：先转 docx 再解析
    for fp in doc_files:
        logger.info(f"转换并解析 .doc 文件: {os.path.basename(fp)}")
        try:
            markdown_content, bbox_map, extra_info = parser.parse_doc(fp, output_dir)
            results[os.path.basename(fp)] = {
                "file": fp,
                "success": True,
            }
        except Exception as e:
            logger.error(f"  ✗ {os.path.basename(fp)}: {e}")
            results[os.path.basename(fp)] = {
                "file": fp,
                "success": False,
                "error": str(e),
            }

    success = sum(1 for r in results.values() if r.get("success"))
    logger.info(f"Word 文档批量解析完成: 成功 {success}/{len(file_paths)}")
    return results
